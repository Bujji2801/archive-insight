from docx import Document
import os

def create_test_docx():
    document = Document()

    document.add_heading('Advanced Traffic Management System', 0)

    p = document.add_paragraph('A comprehensive ')
    p.add_run('Smart Traffic Management System').bold = True
    p.add_run(' designed to optimize and automate traffic flow using ')
    p.add_run('Python and Computer Vision').italic = True
    p.add_run('.')

    document.add_heading('Abstract', level=1)
    document.add_paragraph(
        'This project proposes a framework to revolutionize urban traffic control. '
        'Leveraging modern technologies to enhance real-time analytics and reduce congestion. '
        'The system uses image recognition to monitor vehicle density and adjust signal timings dynamically.'
    )

    document.add_heading('Methodology', level=1)
    document.add_paragraph(
        'We utilize an agile methodology combined with state-of-the-art algorithms like YOLO for object detection. '
        'The platform is built using Python, OpenCV, and a cloud-based dashboard for monitoring.'
    )

    document.add_heading('Expected Outcome', level=1)
    document.add_paragraph(
        'Improvement in efficiency by 40% and reduction in costs associated with manual traffic management.'
    )
    
    # Save to Desktop for easy access
    desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop') 
    # Fallback if USERPROFILE is weird, though it should be fine on Windows.
    # Actually, the user's workspace is c:\Users\ADMIN\OneDrive\Desktop\archive-insight
    # I should save it there for convenience or just "test_project.docx" in the root.
    
    output_path = "test_project.docx"
    document.save(output_path)
    print(f"Created {output_path}")

if __name__ == "__main__":
    create_test_docx()
